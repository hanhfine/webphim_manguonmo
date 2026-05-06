from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SOURCE = Path("/Users/nguyenhoangduong/Documents/webphim_manguonmo/Bao_cao_Phan_mem_ma_nguon_mo_MYCINEMA.docx")
OUTPUT = Path("/Users/nguyenhoangduong/Documents/webphim_manguonmo/Bao_cao_Phan_mem_ma_nguon_mo_MYCINEMA.docx")

CHAPTER_RE = re.compile(r"^CHƯƠNG\s+\d+", re.IGNORECASE)
SECTION_RE = re.compile(r"^\d+(?:\.\d+)*\s")


def set_font(run, size=13, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def configure_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)


def reset_para(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, line=1.15, indent=Cm(1)):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.first_line_indent = indent


def add_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, bold=False, italic=False, before=0, after=6, line=1.15, indent=Cm(1)):
    p = doc.add_paragraph()
    reset_para(p, align=align, before=before, after=after, line=line, indent=indent)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p


def set_cell_text(cell, text, size=13, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    reset_para(p, align=align, before=0, after=0, line=1.0, indent=None)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_border(section):
    clear_page_border(section)
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "24")
        border.set(qn("w:space"), "24")
        border.set(qn("w:color"), "000000")
        pg_borders.append(border)
    section._sectPr.append(pg_borders)


def clear_page_border(section):
    sect_pr = section._sectPr
    for node in sect_pr.findall(qn("w:pgBorders")):
        sect_pr.remove(node)


def parse_source():
    src = Document(str(SOURCE))
    paragraphs = [p.text.strip() for p in src.paragraphs if p.text.strip()]
    start = paragraphs.index("MỞ ĐẦU")
    body = paragraphs[start:]

    intro = []
    chapters = []
    current_chapter = None
    current_item = None
    mode = "intro"

    for text in body:
        if text == "MỞ ĐẦU":
            continue

        if text == "KẾT LUẬN":
            current_chapter = {"title": text, "items": [{"title": "", "paragraphs": []}]}
            chapters.append(current_chapter)
            current_item = current_chapter["items"][0]
            mode = "chapter"
            continue

        if CHAPTER_RE.match(text):
            current_chapter = {"title": text, "items": []}
            chapters.append(current_chapter)
            current_item = None
            mode = "chapter"
            continue

        if SECTION_RE.match(text):
            current_item = {"title": text, "paragraphs": []}
            if current_chapter is None:
                raise RuntimeError(f"Section heading without chapter: {text}")
            current_chapter["items"].append(current_item)
            continue

        if mode == "intro":
            intro.append(text)
        else:
            if current_item is None:
                current_item = {"title": "", "paragraphs": []}
                current_chapter["items"].append(current_item)
            current_item["paragraphs"].append(text)

    return intro, chapters


def build_toc(chapters):
    page = 4
    entries = [("MỞ ĐẦU", str(page), 0, True)]
    for chapter in chapters:
        page += 1
        entries.append((chapter["title"], str(page), 0, True))
        first = True
        for item in chapter["items"]:
            if item["title"]:
                entries.append((item["title"], str(page), 1, False))
            if not first:
                pass
            first = False
            page += 1
        page -= 1
    return entries


def add_cover(doc):
    for _ in range(2):
        doc.add_paragraph()

    add_para(doc, "TRƯỜNG ĐẠI HỌC [TÊN TRƯỜNG]", align=WD_ALIGN_PARAGRAPH.CENTER, size=17, before=0, after=2, line=1.0, indent=None)
    add_para(doc, "KHOA CÔNG NGHỆ THÔNG TIN", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, before=0, after=16, line=1.0, indent=None)
    add_para(doc, "[CHÈN LOGO TRƯỜNG]", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, italic=True, before=8, after=10, line=1.0, indent=None)

    for _ in range(3):
        doc.add_paragraph()

    add_para(doc, "BÁO CÁO MÔN HỌC", align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True, before=0, after=0, line=1.0, indent=None)
    add_para(doc, "PHẦN MỀM MÃ NGUỒN MỞ", align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True, before=0, after=18, line=1.0, indent=None)

    p = doc.add_paragraph()
    reset_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=10, line=1.0, indent=None)
    r = p.add_run("ĐỀ TÀI:")
    set_font(r, size=18, bold=True)
    r.underline = True

    add_para(doc, "XÂY DỰNG HỆ THỐNG ĐẶT VÉ XEM PHIM ONLINE", align=WD_ALIGN_PARAGRAPH.CENTER, size=17, bold=True, before=20, after=28, line=1.0, indent=None)

    table = doc.add_table(rows=5, cols=3)
    table.autofit = False
    widths = [Cm(5.2), Cm(0.8), Cm(8.8)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    rows = [
        ("Sinh viên thực hiện", "[Điền họ tên sinh viên]"),
        ("Giảng viên hướng dẫn", "[Điền tên giảng viên]"),
        ("Ngành", "CÔNG NGHỆ THÔNG TIN"),
        ("Chuyên ngành", "CÔNG NGHỆ PHẦN MỀM"),
        ("Lớp", "[Điền lớp]"),
    ]
    for row, values in zip(table.rows, rows):
        set_cell_text(row.cells[0], values[0], size=14, bold=True)
        set_cell_text(row.cells[1], ":", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], values[1], size=14, bold=(values[0] != "Sinh viên thực hiện"))

    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, "Năm học 2025 - 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, italic=True, before=18, after=0, line=1.0, indent=None)


def add_score_sheet(doc):
    add_para(doc, "PHIẾU CHẤM ĐIỂM", align=WD_ALIGN_PARAGRAPH.LEFT, size=18, bold=True, before=0, after=12, line=1.0, indent=None)

    top = doc.add_table(rows=4, cols=5)
    top.style = "Table Grid"
    top.autofit = False
    widths = [Cm(1.5), Cm(5.4), Cm(8.2), Cm(2.0), Cm(2.0)]
    for row in top.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    headers = ["STT", "Họ và tên sinh viên", "Nội dung thực hiện", "Điểm", "Chữ ký"]
    for idx, header in enumerate(headers):
        set_cell_text(top.rows[0].cells[idx], header, size=15, bold=True)
    for ridx in range(1, 4):
        top.rows[ridx].height = Cm(3.2)
        top.rows[ridx].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        set_cell_text(top.rows[ridx].cells[0], str(ridx), size=15)
        for cidx in range(1, 5):
            set_cell_text(top.rows[ridx].cells[cidx], "", size=13)

    doc.add_paragraph()

    bottom = doc.add_table(rows=3, cols=3)
    bottom.style = "Table Grid"
    bottom.autofit = False
    widths = [Cm(6.5), Cm(6.3), Cm(4.5)]
    for row in bottom.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    headers = ["Họ và tên giảng viên", "Chữ ký", "Ghi chú"]
    for idx, header in enumerate(headers):
        set_cell_text(bottom.rows[0].cells[idx], header, size=15, bold=True)
    labels = ["Giảng viên chấm 1:", "Giảng viên chấm 2:"]
    for ridx, label in enumerate(labels, start=1):
        bottom.rows[ridx].height = Cm(3)
        bottom.rows[ridx].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        set_cell_text(bottom.rows[ridx].cells[0], label, size=15)
        set_cell_text(bottom.rows[ridx].cells[1], "", size=13)
        set_cell_text(bottom.rows[ridx].cells[2], "", size=13)


def add_toc(doc, entries):
    add_para(doc, "MỤC LỤC", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, before=0, after=12, line=1.0, indent=None)
    for text, page, level, strong in entries:
        p = doc.add_paragraph()
        reset_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=3, line=1.0, indent=None)
        p.paragraph_format.left_indent = Cm(0.7 * level)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.8), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r1 = p.add_run(text)
        set_font(r1, size=15 if strong else 14, bold=strong)
        r2 = p.add_run("\t" + page)
        set_font(r2, size=15 if strong else 14, bold=False)


def add_body_footer(section):
    section.footer.is_linked_to_previous = False
    footer = section.footer
    p = footer.paragraphs[0]
    reset_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=1.0, indent=None)
    r = p.add_run("Báo cáo môn học Phần mềm mã nguồn mở")
    set_font(r, size=10, italic=True)


def clear_footer(section):
    section.footer.is_linked_to_previous = False
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    reset_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=1.0, indent=None)


def add_body_section(doc, title, paragraphs):
    if title:
        add_para(doc, title, align=WD_ALIGN_PARAGRAPH.LEFT, size=15, bold=True, before=0, after=8, line=1.0, indent=None)
    for text in paragraphs:
        if text.startswith("[CHÈN HÌNH ẢNH"):
            add_para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True, before=6, after=10, line=1.0, indent=None)
        else:
            add_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, before=0, after=4, line=1.15, indent=Cm(1))


def build():
    intro, chapters = parse_source()
    toc_entries = build_toc(chapters)

    doc = Document()
    configure_page(doc.sections[0])
    add_page_border(doc.sections[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)

    add_cover(doc)
    clear_footer(doc.sections[0])
    add_page_border(doc.sections[0])

    score_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_page(score_section)
    clear_page_border(score_section)
    clear_footer(score_section)
    add_score_sheet(doc)

    toc_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_page(toc_section)
    clear_page_border(toc_section)
    clear_footer(toc_section)
    add_toc(doc, toc_entries)

    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_page(body_section)
    clear_page_border(body_section)
    add_body_footer(body_section)

    add_para(doc, "MỞ ĐẦU", align=WD_ALIGN_PARAGRAPH.CENTER, size=17, bold=True, before=0, after=12, line=1.0, indent=None)
    for text in intro:
        add_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, before=0, after=4, line=1.15, indent=Cm(1))

    for chapter in chapters:
        doc.add_page_break()
        add_para(doc, chapter["title"], align=WD_ALIGN_PARAGRAPH.CENTER, size=17, bold=True, before=0, after=12, line=1.0, indent=None)
        first = True
        for item in chapter["items"]:
            if not first:
                doc.add_page_break()
            add_body_section(doc, item["title"], item["paragraphs"])
            first = False

    doc.save(str(OUTPUT))


if __name__ == "__main__":
    build()
